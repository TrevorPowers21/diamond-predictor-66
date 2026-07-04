#!/usr/bin/env python3
"""Render scripts/.prod_rosters.json (from `npm run export-rosters:prod`) into a
human-readable landscape .docx + two CSVs (rosters, target boards). Coach data —
outputs are gitignored. Usage: python3 scripts/render_roster_docx.py [YYYY_MM_DD]"""
import json, csv, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.section import WD_ORIENT
STAMP = sys.argv[1] if len(sys.argv) > 1 else "export"
data=json.load(open("scripts/.prod_rosters.json"))
NAVY=RGBColor(0x0a,0x14,0x28); GOLD=RGBColor(0xA0,0x88,0x20); GREY=RGBColor(0x55,0x55,0x55); RED=RGBColor(0x99,0x33,0x33)
KS={"Returner":"Returner","Transfer / portal":"Transfer","Imported (hand-entered)":"Imported","Leaving":"Leaving"}
def dol(v): return "—" if v is None else "$"+format(int(v),",")
def dev(p): return "—" if p["devAgg"] is None else f'{p["devAgg"]}'+(" (set)" if p["devAggByCoach"] else "")
def pay(p): a=dol(p["actualPay"]); return a+(" (set)" if p["payByCoach"] and p["actualPay"] is not None else "")
d=Document(); sec=d.sections[0]; sec.orientation=WD_ORIENT.LANDSCAPE
sec.page_width,sec.page_height=Inches(11),Inches(8.5); sec.left_margin=sec.right_margin=Inches(0.5); sec.top_margin=sec.bottom_margin=Inches(0.5)
d.styles['Normal'].font.name='Calibri'; d.styles['Normal'].font.size=Pt(9.5)
def run(p,t,s=9.5,b=False,c=None,i=False):
    r=p.add_run(t); r.font.size=Pt(s); r.bold=b; r.italic=i
    if c: r.font.color.rgb=c
    return r
def head(t,s=14,c=NAVY,bf=10):
    p=d.add_paragraph(); p.space_before=Pt(bf); p.space_after=Pt(2); run(p,t,s,True,c)
p=d.add_paragraph(); run(p,"RSTR IQ — Prod Roster Ground Truth",20,True,NAVY)
p=d.add_paragraph(); run(p,f"Complete human-readable snapshot of every Team Builder roster on production · {STAMP} · READ-ONLY",10,False,GOLD)
p=d.add_paragraph(); run(p,"Purpose: worst-case reference. If any roster looks wrong after deploy, this is the exact 'before' — every player, every on-roster toggle, every dollar — to rebuild or set back by hand.",9.5,False,GREY,True)
head("How to read this — plain English",13,GOLD,8)
for k,v in [("Kind","Returner (already on team) · Transfer (portal target) · Imported (recruit/freshman typed in by hand, not in our DB) · Leaving."),
 ("On Roster?","YES = counts toward the team (lineup/rotation, WAR, budget). NO = on the shared Target Board only, not added to this build."),
 ("Pos","Position or pitcher-role slot."),("Depth Role","Depth-chart slot — e.g. everyday starter, bench, weekend starter, high-leverage reliever."),
 ("Dev Agg","Development-aggressiveness projection knob. '(set)' = coach overrode it."),("Projected $","System's suggested market value."),
 ("Actual Pay $","What the coach assigned. '(set)' = manually entered (real coach input, most important to preserve).")]:
    pa=d.add_paragraph(style='List Bullet'); run(pa,k+": ",9.5,True); run(pa,v,9.5)
COLS=["Player","Kind","On Roster?","Pos","Depth Role","Dev Agg","Projected $","Actual Pay $"]; W=[1.9,0.9,0.9,0.7,1.7,0.8,1.1,1.3]
for t in data["teams"]:
    head(t["name"],13,RED if t["name"].startswith("⚠") else NAVY,14)
    if t["coaches"]: p=d.add_paragraph(); run(p,"Coaches: ",9,True,GREY); run(p,", ".join(t["coaches"]),9,False,GREY)
    if t["teamTargetBoard"]:
        p=d.add_paragraph(); run(p,f'Shared Target Board ({len(t["teamTargetBoard"])} — same for every build): ',9.5,True,GOLD)
        run(p,"  •  ".join(f'{x["name"]} ({x["from"] or "?"}, {x["position"] or "?"})' for x in t["teamTargetBoard"]),9)
    for b in t["builds"]:
        c=b["counts"]; p=d.add_paragraph(); p.space_before=Pt(6)
        run(p,f'  BUILD: {b["name"]}',11,True,NAVY)
        run(p,f'   —   Budget {dol(b["totalBudget"])} · {c["onRoster"]} on roster · {c["targetsOffRoster"]} targets · {c["imported"]} imported',9,False,GREY)
        if not b["players"]: run(d.add_paragraph(),"   (no players)",9,i=True,c=GREY); continue
        tb=d.add_table(rows=1,cols=len(COLS)); tb.style='Light Grid Accent 1'
        for i,h in enumerate(COLS): cell=tb.rows[0].cells[i]; cell.width=Inches(W[i]); run(cell.paragraphs[0],h,8.5,True)
        for pl in b["players"]:
            cs=tb.add_row().cells
            vals=[pl["name"],KS.get(pl["kind"],pl["kind"]),"Yes" if pl["onRoster"] else "No",pl["position"] or "—",pl["depthRole"] or "—",dev(pl),dol(pl["projectedMarketValue"]),pay(pl)]
            for i,v in enumerate(vals):
                cs[i].width=Inches(W[i]); rr=run(cs[i].paragraphs[0],str(v),8.5)
                if i==2 and not pl["onRoster"]: rr.font.color.rgb=GREY
d.save(f"Prod_Roster_Ground_Truth_{STAMP}.docx")
with open(f"Prod_Roster_Ground_Truth_{STAMP}.csv","w",newline="") as f:
    w=csv.writer(f); w.writerow(["Team","Build","Player","Kind","OnRoster","Position","DepthRole","DevAgg","DevAgg_CoachSet","ProjectedMarketValue","ActualPay","ActualPay_CoachSet","ClassTransition","ProjectionTier"])
    for t in data["teams"]:
        for b in t["builds"]:
            for pl in b["players"]:
                w.writerow([t["name"],b["name"],pl["name"],pl["kind"],"Yes" if pl["onRoster"] else "No",pl["position"] or "",pl["depthRole"] or "",pl["devAgg"] if pl["devAgg"] is not None else "",pl["devAggByCoach"],pl["projectedMarketValue"] if pl["projectedMarketValue"] is not None else "",pl["actualPay"] if pl["actualPay"] is not None else "",pl["payByCoach"],pl["classTransition"] or "",pl["projectionTier"] or ""])
with open(f"Prod_Target_Boards_{STAMP}.csv","w",newline="") as f:
    w=csv.writer(f); w.writerow(["Team","Player","FromSchool","Position"])
    for t in data["teams"]:
        for x in t["teamTargetBoard"]: w.writerow([t["name"],x["name"],x["from"] or "",x["position"] or ""])
print("rendered", STAMP)
